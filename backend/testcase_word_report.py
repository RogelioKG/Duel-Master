import ast
import os
import re

from docx import Document


def extract_docstrings_from_file(file_path: str) -> list[tuple[str, str, str]]:
    """Extract docstrings from a given Python file."""
    with open(file_path, encoding="utf-8") as file:
        try:
            tree = ast.parse(file.read(), filename=file_path)
        except SyntaxError:
            return []

    docstrings: list[tuple[str, str, str]] = []
    if module_doc := ast.get_docstring(tree):
        docstrings.append(("Module", file_path, module_doc))

    for node in ast.walk(tree):
        if isinstance(node, ast.FunctionDef | ast.AsyncFunctionDef | ast.ClassDef):
            if docstring := ast.get_docstring(node):
                docstrings.append((node.__class__.__name__, node.name, docstring))

    return docstrings


def collect_docstrings(directory: str) -> list[tuple[str, str, str]]:
    """Recursively collect docstrings from all Python files in the directory."""
    all_docstrings: list[tuple[str, str, str]] = []
    for root, _, files in os.walk(directory):
        for file in files:
            if file.endswith(".py"):
                file_path = os.path.join(root, file)
                all_docstrings.extend(extract_docstrings_from_file(file_path))
    return all_docstrings


def save_docstrings_to_file(docstrings: list[tuple[str, str, str]], output_file: str) -> None:
    """Save collected docstrings to a text file."""
    with open(output_file, mode="w", encoding="UTF-8") as file:
        for item in docstrings:
            file.write(f"{item[0]}: {item[1]}\nDocstring:\n{item[2]}\n{'-' * 40}\n")


def parse_docstrings_from_file(input_file: str) -> list[dict[str, str]]:
    """Extract structured test case information from a docstring file."""
    with open(input_file, encoding="UTF-8") as doc:
        docstring = doc.read()

    pattern = r"""
    FunctionDef:\s*(?P<testcase>\S+)\s*
    Docstring:\s*
    (?P<description>.*?)
    ((M|S)UT\s*---\s*
    .*?\s*
    Description\s*-----------\s*
    \+\s*Given\uFF1A\s*(?P<given>.*?)\s*
    \+\s*When\uFF1A\s*(?P<when>.*?)\s*
    \+\s*Then\uFF1A\s*(?P<then>.*?)\s*
    )?
    ----------------------------------------
    """

    test_cases: list[dict[str, str]] = []
    matches = re.finditer(pattern, docstring, re.VERBOSE | re.DOTALL)

    for match in matches:
        testcase = match.group("testcase")
        description = match.group("description").strip()

        test_case: dict[str, str] = {
            "testcase": f"Testcase{len(test_cases) + 1}",
            "test_function": testcase,
            "description": description,
        }

        given = match.group("given")
        when = match.group("when")
        then = match.group("then")

        if given and when and then:
            test_case.update(
                {
                    "given": given.strip(),
                    "when": when.strip(),
                    "then": then.strip(),
                }
            )

            test_cases.append(test_case)

    return test_cases


def generate_word_document(test_cases: list[dict[str, str]], output_file: str) -> None:
    """Generate a Word document from test cases."""
    word = Document()
    for test_case in test_cases:
        table = word.add_table(rows=5, cols=2)
        table.style = "Table Grid"

        data = [
            ["測試案例編號", test_case.get("testcase", "...")],
            ["測試函數名稱", test_case.get("test_function", "...")],
            ["前置條件 (Given)", test_case.get("given", "...")],
            ["發生情境 (When)", test_case.get("when", "...")],
            ["預期結果 (Then)", test_case.get("then", "...")],
        ]

        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.text = data[i][j]

        word.add_paragraph()

    word.save(output_file)


def main() -> None:
    directory: str = "tests"
    docstring_file: str = "docstring.txt"
    word_file: str = "test_cases.docx"

    docstrings = collect_docstrings(directory)
    save_docstrings_to_file(docstrings, docstring_file)
    test_cases = parse_docstrings_from_file(docstring_file)
    generate_word_document(test_cases, word_file)


if __name__ == "__main__":
    main()
